"""
Resume History MCP — Extractor
================================
Extracts raw text and structured sections from PDF and DOCX files.

Strategy:
  PDF  → pdfplumber (primary) → PyMuPDF (fallback)
  DOCX → python-docx

Mode: LENIENT — extract everything, preserve unclassified sections,
      let downstream (LLM or aggregator) decide what to do with uncertain data.
"""

from __future__ import annotations

import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from .config import (
    IDENTITY_PATTERNS,
    LENIENT_MIN_LINE_LENGTH,
    LENIENT_SECTION_MIN_CONFIDENCE,
    MAX_PAGES_PER_PDF,
    SECTION_KEYWORDS,
)
from .exceptions import DOCXExtractionError, PDFExtractionError
from .schemas import RawResume, RawSection

logger = logging.getLogger(__name__)


# ──────────────────────────────────────────────────────────────────────
# Public Interface
# ──────────────────────────────────────────────────────────────────────


def extract(file_path: Path) -> RawResume:
    """
    Extract a single resume file into a RawResume object.
    Dispatches to the correct extractor based on file extension.
    Always succeeds in lenient mode — failures are stored as warnings.
    """
    suffix = file_path.suffix.lower()
    modified = datetime.fromtimestamp(file_path.stat().st_mtime)

    resume = RawResume(
        file_path=file_path,
        file_name=file_path.name,
        modified_date=modified,
    )

    if suffix == ".pdf":
        _extract_pdf(file_path, resume)
    elif suffix == ".docx":
        _extract_docx(file_path, resume)
    else:
        resume.extraction_warnings.append(f"Unsupported file type: {suffix}")
        return resume

    # Lenient post-processing — always run even if extraction was partial
    _parse_identity(resume)
    _classify_sections(resume)

    logger.info(
        "Extracted '%s' via %s — %d section(s), %d warning(s)",
        file_path.name,
        resume.extractor_used,
        len(resume.raw_sections),
        len(resume.extraction_warnings),
    )
    return resume


# ──────────────────────────────────────────────────────────────────────
# PDF Extraction (pdfplumber → PyMuPDF fallback)
# ──────────────────────────────────────────────────────────────────────


def _extract_pdf(path: Path, resume: RawResume) -> None:
    """Try pdfplumber first; fall back to PyMuPDF on any failure."""
    try:
        _extract_pdf_pdfplumber(path, resume)
        if not resume.full_text.strip():
            raise RuntimeError("pdfplumber returned empty text")
        return
    except Exception as e:
        warn = f"pdfplumber failed on '{path.name}': {e} — trying PyMuPDF"
        logger.warning(warn)
        resume.extraction_warnings.append(warn)

    try:
        _extract_pdf_pymupdf(path, resume)
        if not resume.full_text.strip():
            raise RuntimeError("PyMuPDF returned empty text")
        return
    except Exception as e:
        err = f"PyMuPDF also failed on '{path.name}': {e}"
        logger.error(err)
        resume.extraction_warnings.append(err)
        raise PDFExtractionError(str(path), f"Both pdfplumber and PyMuPDF failed: {e}")


def _extract_pdf_pdfplumber(path: Path, resume: RawResume) -> None:
    import pdfplumber  # lazy import

    lines: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        pages = pdf.pages[:MAX_PAGES_PER_PDF]
        for page in pages:
            text = page.extract_text(x_tolerance=3, y_tolerance=3)
            if text:
                lines.extend(text.splitlines())

    resume.full_text = "\n".join(
        line for line in lines if len(line.strip()) >= LENIENT_MIN_LINE_LENGTH
    )
    resume.extractor_used = "pdfplumber"


def _extract_pdf_pymupdf(path: Path, resume: RawResume) -> None:
    import fitz  # PyMuPDF — lazy import

    lines: list[str] = []
    doc = fitz.open(str(path))
    for page_num, page in enumerate(doc):
        if page_num >= MAX_PAGES_PER_PDF:
            break
        text = page.get_text("text")
        if text:
            lines.extend(text.splitlines())
    doc.close()

    resume.full_text = "\n".join(
        line for line in lines if len(line.strip()) >= LENIENT_MIN_LINE_LENGTH
    )
    resume.extractor_used = "pymupdf"


# ──────────────────────────────────────────────────────────────────────
# DOCX Extraction (python-docx)
# ──────────────────────────────────────────────────────────────────────


def _extract_docx(path: Path, resume: RawResume) -> None:
    try:
        import docx  # python-docx — lazy import

        doc = docx.Document(str(path))
        lines: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if len(text) >= LENIENT_MIN_LINE_LENGTH:
                lines.append(text)

        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    lines.append(row_text)

        resume.full_text = "\n".join(lines)
        resume.extractor_used = "docx"

    except Exception as e:
        err = f"DOCX extraction failed: {e}"
        logger.error(err)
        resume.extraction_warnings.append(err)
        raise DOCXExtractionError(str(path), str(e))


# ──────────────────────────────────────────────────────────────────────
# Identity Parsing (email, phone, links)
# ──────────────────────────────────────────────────────────────────────


def _parse_identity(resume: RawResume) -> None:
    """Extract name and contact info from the top of the full text."""
    if not resume.full_text:
        return

    lines = resume.full_text.splitlines()
    identity: dict[str, Any] = {}

    # Name heuristic: first non-empty line that is not a URL and is title-case
    for line in lines[:8]:
        stripped = line.strip()
        if not stripped:
            continue
        if "@" in stripped or "http" in stripped or re.match(r"^\d", stripped):
            continue
        words = stripped.split()
        if 1 <= len(words) <= 5 and all(not w[0].islower() for w in words if w):
            identity["name"] = stripped
            break

    # Contact fields
    full = resume.full_text
    for field, pattern in IDENTITY_PATTERNS.items():
        match = pattern.search(full)
        if match:
            identity[field] = match.group(0)

    resume.identity = identity


# ──────────────────────────────────────────────────────────────────────
# Section Classification (lenient — capture everything)
# ──────────────────────────────────────────────────────────────────────


def _classify_sections(resume: RawResume) -> None:
    """
    Split full_text into sections by detecting heading lines.
    LENIENT MODE: every detected block is stored, confidence may be 0.
    Unclassified blocks are preserved in raw_sections for LLM use.
    """
    if not resume.full_text:
        return

    lines = resume.full_text.splitlines()
    sections: list[tuple[int, str, str]] = []  # (line_idx, section_name, heading_text)

    # Build reverse lookup: keyword → canonical section name
    keyword_map: dict[str, str] = {}
    for section_name, keywords in SECTION_KEYWORDS.items():
        for kw in keywords:
            keyword_map[kw.lower()] = section_name

    for idx, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            continue
        normalized = re.sub(r"[^a-z\s]", "", stripped.lower()).strip()

        # Direct keyword match
        for kw, section_name in keyword_map.items():
            if normalized == kw or normalized.startswith(kw + " "):
                sections.append((idx, section_name, stripped))
                break
        else:
            # Heuristic: short all-caps line is likely a heading (lenient)
            if (
                stripped.isupper()
                and 3 <= len(stripped) <= 40
                and len(stripped.split()) <= 5
            ):
                sections.append((idx, "unclassified", stripped))

    if not sections:
        # No sections found — store entire text as one unclassified block
        resume.raw_sections.append(
            RawSection(
                heading="FULL TEXT",
                content=resume.full_text,
                confidence=0.0,
            )
        )
        return

    # Slice lines between headings into section content
    for i, (start_idx, section_name, heading_text) in enumerate(sections):
        end_idx = sections[i + 1][0] if i + 1 < len(sections) else len(lines)
        content = "\n".join(
            line for line in lines[start_idx + 1 : end_idx] if line.strip()
        ).strip()

        confidence = 0.0 if section_name == "unclassified" else 0.8
        raw_sec = RawSection(
            heading=heading_text,
            content=content,
            confidence=confidence,
        )

        if confidence >= LENIENT_SECTION_MIN_CONFIDENCE:
            resume.raw_sections.append(raw_sec)

        # Populate typed fields for confident matches
        if section_name == "summary":
            resume.summary = resume.summary or content
        elif section_name == "skills":
            resume.skills.extend(_parse_skills(content))
        elif section_name == "certifications":
            resume.certifications.extend(_parse_bullet_list(content))
        elif section_name == "experience":
            if content:
                resume.experiences.append({"raw": content})
        elif section_name == "projects":
            if content:
                resume.projects.append({"raw": content})
        elif section_name == "education":
            if content:
                resume.education.append({"raw": content})


# ──────────────────────────────────────────────────────────────────────
# Micro-parsers
# ──────────────────────────────────────────────────────────────────────


def _parse_skills(content: str) -> list[str]:
    """Extract individual skill tokens from a skills section."""
    skills: list[str] = []
    for line in content.splitlines():
        # Split by common delimiters
        tokens = re.split(r"[,|•·\t]+", line)
        for token in tokens:
            cleaned = token.strip().strip(":-")
            if 1 < len(cleaned) <= 50:
                skills.append(cleaned)
    return skills


def _parse_bullet_list(content: str) -> list[str]:
    """Extract bullet/line items as a flat list."""
    items: list[str] = []
    for line in content.splitlines():
        cleaned = re.sub(r"^[\s•·\-*]+", "", line).strip()
        if len(cleaned) > 3:
            items.append(cleaned)
    return items
